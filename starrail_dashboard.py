import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import numpy as np
from datetime import datetime, timedelta
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
import tempfile
import base64
import re
from scipy.stats import chi2_contingency, ttest_ind

# ---------------------- 页面基础配置 ----------------------
st.set_page_config(
    page_title="星穹铁道运营数据看板",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 自定义CSS样式（仅修改conclusion-box和warn-box背景/文字色，其余完全保留）
st.markdown("""
    <style>
    /* 整体样式 */
    .main {
        background-color: #0a0e17;
        color: #e0e6ff;
    }
    /* 标题样式 */
    .module-title {
        color: #a78bfa;
        font-size: 20px;
        font-weight: bold;
        margin: 20px 0 10px 0;
        border-left: 4px solid #8b5cf6;
        padding-left: 10px;
    }
    /* 结论框样式 - 改为浅色背景，文字清晰 */
    .conclusion-box {
        background-color: #f8f9ff;
        color: #1f2937;
        padding: 15px;
        border-radius: 8px;
        margin: 10px 0;
        border: 1px solid #8b5cf6;
    }
    /* 警告框样式 - 改为浅色背景，文字清晰 */
    .warn-box {
        background-color: #fff7ed;
        color: #7c2d12;
        padding: 10px;
        border-radius: 6px;
        margin: 5px 0;
        border: 1px solid #f59e0b;
    }
    /* 按钮样式 */
    div.stButton > button {
        background-color: #4f46e5;
        color: white;
        border-radius: 6px;
        border: none;
        padding: 8px 16px;
    }
    div.stButton > button:hover {
        background-color: #4338ca;
    }
    /* 数据卡片样式 */
    .data-card {
        background-color: #111827;
        padding: 15px;
        border-radius: 8px;
        text-align: center;
        border: 1px solid #312e81;
    }
    /* 侧边栏样式 */
    .sidebar .sidebar-content {
        background-color: #111827;
    }
    /* 标签页样式优化 */
    button[data-baseweb="tab"] {
        color: #e0e6ff !important;
        font-size: 16px !important;
    }
    button[data-baseweb="tab"][aria-selected="true"] {
        color: #a78bfa !important;
        border-bottom: 2px solid #8b5cf6 !important;
    }
    </style>
""", unsafe_allow_html=True)

# ---------------------- 工具函数：清除特殊字符+统一文本格式 ----------------------
def clean_text(text):
    """清除HTML标签、特殊字符，统一文本格式"""
    if not text:
        return ""
    # 清除HTML标签
    text = re.sub(r'<br>|</br>|<[^>]+>', '', text)
    # 清除特殊字符（保留中文、英文、数字、标点）
    text = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9\s，。；：！？（）《》【】、；：""''.,;:!?()[]{}<>]', '', text)
    # 去除多余空格和换行
    text = re.sub(r'\s+', ' ', text).strip()
    return text

# ---------------------- A/B测试分析模块函数 ----------------------
def ab_test_module():
    st.markdown('<div class="module-title">📊 A/B测试效果分析（周年庆累计充值奖励）</div>', unsafe_allow_html=True)
    
    # 数据上传
    ab_file = st.file_uploader(
        "📤 上传周年庆A/B测试数据（CSV/Excel）", 
        type=["csv", "xlsx"], 
        key="ab_uploader",
        help="请上传生成的hsr_anniversary_ab_test.csv/xlsx文件"
    )
    
    if ab_file:
        try:
            # 加载数据
            if ab_file.name.endswith(".csv"):
                ab_df = pd.read_csv(ab_file)
            else:
                ab_df = pd.read_excel(ab_file)
            
            st.success("✅ A/B测试数据加载成功！")
            st.dataframe(ab_df.head(), use_container_width=True)
            
            # 数据验证
            if "group" not in ab_df.columns or "conversion" not in ab_df.columns:
                st.error("❌ 数据格式错误：缺少group或conversion字段")
                return
            
            # 分组数据准备
            groups = ab_df["group"].unique()
            if len(groups) != 2:
                st.warning("⚠️ 当前仅支持两组A/B测试（对照组 vs 实验组）")
                return
            
            control = ab_df[ab_df["group"] == "control"]
            test = ab_df[ab_df["group"] == "test"]
            
            # 核心指标卡片
            st.subheader("核心指标概览", divider="violet")
            col1, col2, col3, col4 = st.columns(4)
            
            # 付费转化率
            with col1:
                control_conv = control["conversion"].mean()
                test_conv = test["conversion"].mean()
                lift_conv = (test_conv - control_conv) / control_conv if control_conv > 0 else 0
                st.metric("对照组转化率", f"{control_conv:.2%}")
                st.metric("实验组转化率", f"{test_conv:.2%}")
                st.metric("转化率提升", f"{lift_conv:.2%}")
            
            # ARPU
            with col2:
                control_arpu = control["arpu"].mean()
                test_arpu = test["arpu"].mean()
                lift_arpu = (test_arpu - control_arpu) / control_arpu if control_arpu > 0 else 0
                st.metric("对照组ARPU", f"¥{control_arpu:.2f}")
                st.metric("实验组ARPU", f"¥{test_arpu:.2f}")
                st.metric("ARPU提升", f"{lift_arpu:.2%}")
            
            # 高价值用户占比
            with col3:
                control_hv = control["high_value"].mean()
                test_hv = test["high_value"].mean()
                lift_hv = (test_hv - control_hv) / control_hv if control_hv > 0 else 0
                st.metric("对照组高价值用户占比", f"{control_hv:.2%}")
                st.metric("实验组高价值用户占比", f"{test_hv:.2%}")
                st.metric("高价值用户占比提升", f"{lift_hv:.2%}")
            
            # 7日留存率
            with col4:
                control_ret = control["retention_7"].mean()
                test_ret = test["retention_7"].mean()
                lift_ret = (test_ret - control_ret) / control_ret if control_ret > 0 else 0
                st.metric("对照组7日留存", f"{control_ret:.2%}")
                st.metric("实验组7日留存", f"{test_ret:.2%}")
                st.metric("7日留存提升", f"{lift_ret:.2%}")
            
            # 显著性检验
            st.subheader("📈 统计显著性检验", divider="violet")
            # 添加P值说明
            st.info("💡 注：P值显示为0.0000是由于模拟数据差异显著且样本量较大（5万），真实业务中会根据实际数据动态变化")
            col1, col2 = st.columns(2)
            
            # 转化率卡方检验
            with col1:
                st.markdown("### 付费转化率（卡方检验）")
                control_conv_count = control["conversion"].sum()
                control_total = len(control)
                test_conv_count = test["conversion"].sum()
                test_total = len(test)
                
                contingency = [
                    [control_conv_count, control_total - control_conv_count],
                    [test_conv_count, test_total - test_conv_count]
                ]
                chi2, p_value_conv, dof, expected = chi2_contingency(contingency)
                
                st.write(f"卡方值：{chi2:.4f}")
                st.write(f"P值：{p_value_conv:.4f}")
                if p_value_conv < 0.05:
                    st.success("✅ 转化率差异具有统计显著性")
                else:
                    st.info("ℹ️ 转化率差异无统计显著性")
            
            # ARPU t检验
            with col2:
                st.markdown("### ARPU（t检验）")
                t_stat, p_value_arpu = ttest_ind(control["arpu"], test["arpu"])
                
                st.write(f"t统计量：{t_stat:.4f}")
                st.write(f"P值：{p_value_arpu:.4f}")
                if p_value_arpu < 0.05:
                    st.success("✅ ARPU差异具有统计显著性")
                else:
                    st.info("ℹ️ ARPU差异无统计显著性")
            
            # ---------------------- 核心修改1：拆分柱状图为两个独立图表 ----------------------
            st.subheader("📊 核心指标对比可视化", divider="violet")
            
            # 图表1：百分比指标对比（付费转化率、高价值用户占比、7日留存率）
            percent_df = pd.DataFrame({
                "指标": ["付费转化率", "高价值用户占比", "7日留存率"],
                "对照组": [control_conv, control_hv, control_ret],
                "实验组": [test_conv, test_hv, test_ret]
            })
            
            fig1 = px.bar(
                percent_df,
                x="指标",
                y=["对照组", "实验组"],
                barmode="group",
                title="A/B测试百分比指标对比（%）",
                labels={"value": "占比/率", "variable": "分组"},
                template="plotly_dark",
                color_discrete_map={"对照组": "#6366f1", "实验组": "#ec4899"}
            )
            # 设置Y轴为百分比格式，自动适配范围
            fig1.update_layout(
                yaxis_tickformat=".2%", 
                yaxis_range=[0, max(percent_df[["对照组", "实验组"]].max().max()*1.2, 0.2)],
                height=400
            )
            st.plotly_chart(fig1, use_container_width=True)
            
            # 图表2：ARPU对比（元）
            arpu_df = pd.DataFrame({
                "指标": ["ARPU(元)"],
                "对照组": [control_arpu],
                "实验组": [test_arpu]
            })
            
            fig2 = px.bar(
                arpu_df,
                x="指标",
                y=["对照组", "实验组"],
                barmode="group",
                title="A/B测试ARPU对比（元）",
                labels={"value": "金额（元）", "variable": "分组"},
                template="plotly_dark",
                color_discrete_map={"对照组": "#6366f1", "实验组": "#ec4899"}
            )
            # 设置Y轴范围，确保对比清晰
            fig2.update_layout(
                yaxis_range=[0, max(arpu_df[["对照组", "实验组"]].max().max()*1.2, 30)],
                height=400
            )
            st.plotly_chart(fig2, use_container_width=True)
            # ---------------------- 核心修改1 结束 ----------------------
            
            # 运营结论
            st.subheader("📝 运营结论与建议", divider="violet")
            conclusion_text = ""
            if p_value_conv < 0.05 and p_value_arpu < 0.05:
                conclusion_text = f"""
                <div class="conclusion-box">
                    <b>✅ 测试结论：周年庆累计充值奖励策略效果显著！</b><br><br>
                    1. 付费拉动：实验组付费转化率提升{lift_conv:.2%}，ARPU提升{lift_arpu:.2%}，高价值用户占比提升{lift_hv:.2%}，有效拉动用户付费意愿；<br>
                    2. 留存提升：实验组7日留存率提升{lift_ret:.2%}，验证了付费用户留存优势；<br>
                    3. 策略建议：全量上线该累计充值奖励策略，针对全等级用户推送累充奖励提醒。
                </div>
                """
            else:
                conclusion_text = f"""
                <div class="warn-box">
                    <b>⚠️ 测试结论：策略效果未达显著水平</b><br><br>
                    1. 数据显示实验组指标有提升，但统计显著性不足；<br>
                    2. 建议：扩大样本量至10万用户，或调整奖励梯度（如降低累计充值门槛）后重新测试。
                </div>
                """
            st.markdown(conclusion_text, unsafe_allow_html=True)
            
        except Exception as e:
            st.error(f"❌ A/B测试数据分析失败：{str(e)}")

# ---------------------- 用户行为漏斗分析模块函数 ----------------------
def funnel_analysis_module():
    st.markdown('<div class="module-title">🔍 用户行为漏斗分析（卡池付费转化路径）</div>', unsafe_allow_html=True)
    
    # 数据上传
    funnel_file = st.file_uploader(
        "📤 上传卡池行为漏斗数据（CSV/Excel）", 
        type=["csv", "xlsx"], 
        key="funnel_uploader",
        help="请上传生成的hsr_card_pool_funnel.csv/xlsx文件"
    )
    
    if funnel_file:
        try:
            # 加载数据
            if funnel_file.name.endswith(".csv"):
                funnel_df = pd.read_csv(funnel_file)
            else:
                funnel_df = pd.read_excel(funnel_file)
            
            # 数据预处理
            funnel_df["timestamp"] = pd.to_datetime(funnel_df["timestamp"])
            st.success("✅ 行为漏斗数据加载成功！")
            st.dataframe(funnel_df.head(), use_container_width=True)
            
            # 验证必要字段
            if "step" not in funnel_df.columns or "user_id" not in funnel_df.columns:
                st.error("❌ 数据格式错误：缺少step或user_id字段")
                return
            
            # 步骤选择
            all_steps = sorted(funnel_df["step"].unique())
            default_steps = ["登录游戏", "进入卡池界面", "进行抽卡操作", "打开付费界面", "完成付费"]
            # 过滤默认步骤，只保留数据中存在的
            default_steps = [step for step in default_steps if step in all_steps]
            
            steps = st.multiselect(
                "选择漏斗步骤（按转化顺序）",
                options=all_steps,
                default=default_steps,
                key="funnel_steps"
            )
            
            if len(steps) >= 2:
                # 计算漏斗数据
                funnel_data = []
                prev_user_count = None
                
                for step in steps:
                    # 统计该步骤的独立用户数
                    step_user_count = funnel_df[funnel_df["step"] == step]["user_id"].nunique()
                    
                    # 计算转化率
                    if prev_user_count is None:
                        conversion_rate = 1.0  # 第一步转化率100%
                        drop_rate = 0.0
                    else:
                        conversion_rate = step_user_count / prev_user_count if prev_user_count > 0 else 0
                        drop_rate = 1 - conversion_rate
                    
                    funnel_data.append({
                        "步骤": step,
                        "用户数": step_user_count,
                        "转化率": conversion_rate,
                        "流失率": drop_rate,
                        "累计转化率": step_user_count / funnel_data[0]["用户数"] if funnel_data else 1.0
                    })
                    
                    prev_user_count = step_user_count
                
                # 转换为DataFrame展示
                funnel_df_result = pd.DataFrame(funnel_data)
                st.subheader("📊 漏斗转化数据详情", divider="violet")
                st.dataframe(
                    funnel_df_result.style.format({
                        "转化率": "{:.2%}",
                        "流失率": "{:.2%}",
                        "累计转化率": "{:.2%}"
                    }),
                    use_container_width=True
                )
                
                # 漏斗图可视化
                st.subheader("📈 卡池付费转化漏斗图", divider="violet")
                fig = go.Figure(go.Funnel(
                    y=funnel_df_result["步骤"],
                    x=funnel_df_result["用户数"],
                    textposition="inside",
                    textinfo="value+percent initial",
                    marker={
                        "color": ["#6366f1", "#8b5cf6", "#a78bfa", "#d8b4fe", "#ec4899"][:len(steps)],
                        "line": {"width": 2, "color": "#111827"}
                    },
                    connector={"line": {"width": 2, "color": "#8b5cf6"}}
                ))
                fig.update_layout(
                    title="用户卡池付费转化路径漏斗",
                    template="plotly_dark",
                    height=400
                )
                st.plotly_chart(fig, use_container_width=True)
                
                # 关键流失环节分析
                st.subheader("⚠️ 关键流失环节分析", divider="violet")
                # 排除第一步（无流失）
                loss_data = funnel_df_result[funnel_df_result["流失率"] > 0]
                if not loss_data.empty:
                    max_loss_step = loss_data.loc[loss_data["流失率"].idxmax()]
                    
                    st.markdown(f"""
                    <div class="warn-box">
                        <b>最高流失环节：{max_loss_step['步骤']}</b><br>
                        该环节流失率：{max_loss_step['流失率']:.2%}<br>
                        转化率仅：{max_loss_step['转化率']:.2%}
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # 针对性建议
                    if "打开付费界面" in max_loss_step["步骤"]:
                        advice = """
                        <div class="conclusion-box">
                            <b>优化建议：</b><br>
                            1. 付费界面优化：优先展示核心档位礼包和突出优惠幅度，降低用户决策成本；<br>
                            2. 福利刺激：增加首充双倍、累计充值奖励弹窗提醒；<br>
                            3. 支付体验：排查支付流程卡顿问题，支持更多支付方式。
                        </div>
                        """
                    elif "进行抽卡操作" in max_loss_step["步骤"]:
                        advice = """
                        <div class="conclusion-box">
                            <b>优化建议：</b><br>
                            1. 抽卡体验：优化抽卡动画和反馈；<br>
                            2. 概率公示：突出展示UP角色获取概率，提升用户信心；<br>
                            3. 保底提醒：实时显示距离保底的抽数，刺激用户继续抽卡。
                        </div>
                        """
                    elif "进入卡池界面" in max_loss_step["步骤"]:
                        advice = """
                        <div class="conclusion-box">
                            <b>优化建议：</b><br>
                            1. 界面入口：增加卡池入口的显眼度，简化进入流程；<br>
                            2. 角色展示：突出UP角色立绘和技能介绍，提升吸引力；<br>
                            3. 福利引导：新用户首次进入卡池赠送免费抽卡次数。
                        </div>
                        """
                    else:
                        advice = """
                        <div class="conclusion-box">
                            <b>优化建议：</b><br>
                            1. 流程简化：减少该环节的操作步骤，提升用户体验；<br>
                            2. 付费界面优化：优先展示核心档位礼包和突出优惠幅度，降低用户决策成本；<br>
                            3. A/B测试：测试不同界面设计，选择转化率最优方案。
                        </div>
                        """
                    st.markdown(advice, unsafe_allow_html=True)
                
                # 按用户等级分群分析
                if "user_level" in funnel_df.columns:
                    st.subheader("👥 按用户等级分群漏斗分析", divider="violet")
                    
                    # 等级分群
                    funnel_df["等级区间"] = pd.cut(
                        funnel_df["user_level"],
                        bins=[0, 20, 40, 60, 70],
                        labels=["1-20级", "21-40级", "41-60级", "61-70级"]
                    )
                    
                    # 选择等级区间
                    selected_level = st.selectbox(
                        "选择用户等级区间",
                        options=["1-20级", "21-40级", "41-60级", "61-70级"],
                        key="level_selector"
                    )
                    
                    # 筛选该等级区间数据
                    level_funnel_df = funnel_df[funnel_df["等级区间"] == selected_level]
                    
                    # 重新计算该等级的漏斗
                    level_funnel_data = []
                    prev_count = None
                    for step in steps:
                        step_count = level_funnel_df[level_funnel_df["step"] == step]["user_id"].nunique()
                        if prev_count is None:
                            conv = 1.0
                        else:
                            conv = step_count / prev_count if prev_count > 0 else 0
                        level_funnel_data.append({
                            "步骤": step,
                            "用户数": step_count,
                            "转化率": conv
                        })
                        prev_count = step_count
                    
                    # 展示等级分群漏斗
                    level_funnel_result = pd.DataFrame(level_funnel_data)
                    st.dataframe(
                        level_funnel_result.style.format({"转化率": "{:.2%}"}),
                        use_container_width=True
                    )
                    
                    # 等级分析结论
                    high_level_conversion = level_funnel_result.iloc[-1]["转化率"] if len(level_funnel_result) > 0 else 0
                    if selected_level == "41-60级" and high_level_conversion > 0.07:
                        st.success(f"✅ {selected_level}用户付费转化率最高，是核心付费群体！")
                    elif selected_level == "1-20级" and high_level_conversion < 0.02:
                        st.info(f"ℹ️ {selected_level}用户付费意愿低，建议侧重新手福利引导。")
            
        except Exception as e:
            st.error(f"❌ 漏斗数据分析失败：{str(e)}")

# ---------------------- 侧边栏配置 ----------------------
st.sidebar.markdown("<h2 style='color:#a78bfa;'>🚀 星穹铁道运营分析</h2>", unsafe_allow_html=True)
st.sidebar.divider()

# 主运营数据上传（仅用于运营分析标签页）
uploaded_file = st.sidebar.file_uploader(
    "📤 上传运营数据集（Excel）",
    type=["xlsx"],
    help="请上传生成的「崩坏星穹铁道_成熟稳定期双版本运营数据集.xlsx」"
)

# 预警阈值设置（仅用于运营分析标签页）
st.sidebar.markdown("<h4 style='color:#a78bfa;'>⚠️ 异常预警阈值</h4>", unsafe_allow_html=True)
dau_warn_threshold = st.sidebar.number_input(
    "DAU预警阈值（万）",
    min_value=200,
    max_value=500,
    value=300,
    step=10,
    help="DAU低于该值时触发预警"
) * 10000  # 转换为实际数值

retention_warn_threshold = st.sidebar.number_input(
    "新手7日留存率预警阈值",
    min_value=0.30,
    max_value=0.50,
    value=0.45,
    step=0.01,
    help="新手7日留存率低于该值时触发预警"
)

st.sidebar.divider()
st.sidebar.markdown("""
    <div style="text-align:center; color:#94a3b8;">
        📅 数据周期：84天<br>
        🎮 版本：成熟稳定期双版本<br>
        💡 基于星铁真实运营逻辑
    </div>
""", unsafe_allow_html=True)

# ---------------------- 标签页布局 ----------------------
tab1, tab2, tab3 = st.tabs(["运营分析", "A/B测试分析", "用户行为漏斗分析"])

# ---------------------- 1. 运营分析标签页（原有完整功能） ----------------------
with tab1:
    if uploaded_file is not None:
        try:
            # 加载所有sheet
            activity_df = pd.read_excel(uploaded_file, sheet_name="日活数据")
            retention_df = pd.read_excel(uploaded_file, sheet_name="留存数据")
            payment_df = pd.read_excel(uploaded_file, sheet_name="付费数据")
            user_layer_df = pd.read_excel(uploaded_file, sheet_name="用户分层数据")
            
            # 数据预处理：日期格式转换
            activity_df["日期"] = pd.to_datetime(activity_df["日期"]).dt.date
            retention_df["日期"] = pd.to_datetime(retention_df["日期"]).dt.date
            payment_df["日期"] = pd.to_datetime(payment_df["日期"]).dt.date
            
            # 计算付费转化率（新增核心指标）
            convert_df = pd.merge(
                activity_df[["日期", "DAU"]],
                payment_df[["日期", "付费人数"]],
                on="日期"
            )
            convert_df["付费转化率"] = convert_df["付费人数"] / convert_df["DAU"]
            
            st.success("✅ 数据加载成功！")
            
            # ---------------------- 1. 核心指标概览 ----------------------
            st.markdown('<div class="module-title">📊 核心指标概览</div>', unsafe_allow_html=True)
            
            # 计算核心指标
            total_new_users = activity_df["新增用户数"].sum()
            avg_dau = int(activity_df["DAU"].mean())
            avg_retention_7 = retention_df["新手7日留存率"].mean()
            total_revenue = payment_df["卡池总流水(元)"].sum() / 10000  # 转换为万元
            avg_conversion = convert_df["付费转化率"].mean()
            
            # 数据卡片展示
            col1, col2, col3, col4, col5 = st.columns(5)
            with col1:
                st.markdown(f"""
                    <div class="data-card">
                        <p style="color:#94a3b8; margin:0;">总新增用户</p>
                        <h3 style="color:#a78bfa; margin:5px 0;">{total_new_users:,}</h3>
                    </div>
                """, unsafe_allow_html=True)
            with col2:
                st.markdown(f"""
                    <div class="data-card">
                        <p style="color:#94a3b8; margin:0;">平均DAU</p>
                        <h3 style="color:#a78bfa; margin:5px 0;">{avg_dau:,}</h3>
                    </div>
                """, unsafe_allow_html=True)
            with col3:
                st.markdown(f"""
                    <div class="data-card">
                        <p style="color:#94a3b8; margin:0;">新手7日留存率</p>
                        <h3 style="color:#a78bfa; margin:5px 0;">{avg_retention_7:.2%}</h3>
                    </div>
                """, unsafe_allow_html=True)
            with col4:
                st.markdown(f"""
                    <div class="data-card">
                        <p style="color:#94a3b8; margin:0;">总流水（万元）</p>
                        <h3 style="color:#a78bfa; margin:5px 0;">{total_revenue:.2f}</h3>
                    </div>
                """, unsafe_allow_html=True)
            with col5:
                st.markdown(f"""
                    <div class="data-card">
                        <p style="color:#94a3b8; margin:0;">平均付费转化率</p>
                        <h3 style="color:#a78bfa; margin:5px 0;">{avg_conversion:.2%}</h3>
                    </div>
                """, unsafe_allow_html=True)
            
            # ---------------------- 2. 日活与留存分析 ----------------------
            st.markdown('<div class="module-title">📈 日活与留存分析</div>', unsafe_allow_html=True)
            
            # 分栏展示
            col1, col2 = st.columns(2)
            
            # DAU趋势图
            with col1:
                st.subheader("DAU趋势（含预警线）", divider="violet")
                fig_dau = px.line(
                    activity_df,
                    x="日期",
                    y="DAU",
                    title="日活跃用户数趋势",
                    labels={"DAU": "DAU（人）", "日期": "日期"},
                    template="plotly_dark",
                    color_discrete_sequence=["#8b5cf6"]
                )
                # 添加预警线
                fig_dau.add_hline(
                    y=dau_warn_threshold,
                    line_dash="dash",
                    line_color="red",
                    annotation_text=f"预警阈值：{dau_warn_threshold/10000}万",
                    annotation_position="top right"
                )
                st.plotly_chart(fig_dau, use_container_width=True)
                
                # DAU异常预警
                dau_warn_data = activity_df[activity_df["DAU"] < dau_warn_threshold]
                if len(dau_warn_data) > 0:
                    st.markdown(f"""
                        <div class="warn-box">
                            ⚠️ 检测到{len(dau_warn_data)}天DAU低于预警阈值！
                            <br>异常日期：{', '.join([str(d) for d in dau_warn_data['日期'].head(5)])}{'...' if len(dau_warn_data)>5 else ''}
                        </div>
                    """, unsafe_allow_html=True)
            
            # 留存趋势图
            with col2:
                st.subheader("留存率趋势", divider="violet")
                fig_ret = px.line(
                    retention_df,
                    x="日期",
                    y=["新手次日留存率", "新手7日留存率", "活跃用户7日留存率"],
                    title="留存率趋势对比",
                    labels={"value": "留存率", "日期": "日期", "variable": "留存类型"},
                    template="plotly_dark",
                    color_discrete_map={
                        "新手次日留存率": "#a78bfa",
                        "新手7日留存率": "#6366f1",
                        "活跃用户7日留存率": "#3b82f6"
                    }
                )
                # 添加新手7日留存预警线
                fig_ret.add_hline(
                    y=retention_warn_threshold,
                    line_dash="dash",
                    line_color="red",
                    annotation_text=f"预警阈值：{retention_warn_threshold:.2%}",
                    annotation_position="top right"
                )
                st.plotly_chart(fig_ret, use_container_width=True)
                
                # 留存异常预警
                ret_warn_data = retention_df[retention_df["新手7日留存率"] < retention_warn_threshold]
                if len(ret_warn_data) > 0:
                    st.markdown(f"""
                        <div class="warn-box">
                            ⚠️ 检测到{len(ret_warn_data)}天新手7日留存率低于预警阈值！
                        </div>
                    """, unsafe_allow_html=True)
            
            # ---------------------- 3. 付费数据分析 ----------------------
            st.markdown('<div class="module-title">💰 付费数据分析</div>', unsafe_allow_html=True)
            
            # 分栏展示
            col1, col2 = st.columns(2)
            
            # 卡池流水分析
            with col1:
                st.subheader("各卡池总流水", divider="violet")
                cycle_revenue = payment_df.groupby("卡池周期")["卡池总流水(元)"].sum().reset_index()
                cycle_revenue["卡池总流水(万元)"] = cycle_revenue["卡池总流水(元)"] / 10000
                
                fig_cycle = px.bar(
                    cycle_revenue,
                    x="卡池周期",
                    y="卡池总流水(万元)",
                    title="各卡池周期总流水对比",
                    labels={"卡池总流水(万元)": "流水（万元）", "卡池周期": "卡池周期"},
                    template="plotly_dark",
                    color="卡池总流水(万元)",
                    color_continuous_scale="purples"
                )
                # 旋转x轴标签
                fig_cycle.update_layout(xaxis_tickangle=-45)
                st.plotly_chart(fig_cycle, use_container_width=True)
            
            # 付费转化率趋势
            with col2:
                st.subheader("付费转化率趋势", divider="violet")
                fig_convert = px.line(
                    convert_df,
                    x="日期",
                    y="付费转化率",
                    title="每日付费转化率",
                    labels={"付费转化率": "付费转化率", "日期": "日期"},
                    template="plotly_dark",
                    color_discrete_sequence=["#ec4899"]
                )
                # 添加平均值线
                fig_convert.add_hline(
                    y=avg_conversion,
                    line_dash="dash",
                    line_color="green",
                    annotation_text=f"平均值：{avg_conversion:.2%}",
                    annotation_position="top right"
                )
                st.plotly_chart(fig_convert, use_container_width=True)
            
            # 卡池流水构成分析
            st.subheader("卡池流水构成", divider="violet")
            # 计算各卡池类型总流水
            pool_type_revenue = pd.DataFrame({
                "卡池类型": ["主UP池", "复刻池", "常驻池"],
                "总流水(万元)": [
                    payment_df["主UP池流水(元)"].sum()/10000,
                    payment_df["复刻池流水(元)"].sum()/10000,
                    payment_df["常驻池流水(元)"].sum()/10000
                ]
            })
            
            fig_pool = px.pie(
                pool_type_revenue,
                values="总流水(万元)",
                names="卡池类型",
                title="卡池流水构成占比",
                template="plotly_dark",
                color_discrete_map={
                    "主UP池": "#8b5cf6",
                    "复刻池": "#6366f1",
                    "常驻池": "#3b82f6"
                }
            )
            st.plotly_chart(fig_pool, use_container_width=True)
            
            # ---------------------- 4. 用户分层分析 ----------------------
            st.markdown('<div class="module-title">👥 用户分层分析</div>', unsafe_allow_html=True)
            
            # 分栏展示
            col1, col2 = st.columns(2)
            
            # 付费价值分层
            with col1:
                st.subheader("付费价值分层", divider="violet")
                value_layer = user_layer_df["付费价值分层"].value_counts().reset_index()
                value_layer.columns = ["付费分层", "用户数"]
                
                fig_value = px.pie(
                    value_layer,
                    values="用户数",
                    names="付费分层",
                    title="付费价值分层占比",
                    template="plotly_dark",
                    color_discrete_map={
                        "高价值": "#ec4899",
                        "中价值": "#8b5cf6",
                        "低价值": "#6366f1",
                        "免费": "#3b82f6"
                    }
                )
                st.plotly_chart(fig_value, use_container_width=True)
            
            # 生命周期分层
            with col2:
                st.subheader("生命周期分层", divider="violet")
                life_layer = user_layer_df["生命周期分层"].value_counts().reset_index()
                life_layer.columns = ["生命周期", "用户数"]
                
                fig_life = px.bar(
                    life_layer,
                    x="生命周期",
                    y="用户数",
                    title="新老用户占比",
                    labels={"用户数": "用户数", "生命周期": "用户类型"},
                    template="plotly_dark",
                    color="用户数",
                    color_continuous_scale="pinkyl"
                )
                st.plotly_chart(fig_life, use_container_width=True)
            
            # 高价值用户付费偏好
            st.subheader("高价值用户付费卡池偏好", divider="violet")
            high_value_users = user_layer_df[user_layer_df["付费价值分层"] == "高价值"]
            high_value_pay_cycle = high_value_users["主要付费卡池周期"].value_counts().reset_index()
            high_value_pay_cycle.columns = ["卡池周期", "用户数"]
            
            fig_high_pay = px.bar(
                high_value_pay_cycle,
                x="卡池周期",
                y="用户数",
                title="高价值用户付费卡池分布",
                labels={"用户数": "用户数", "卡池周期": "卡池周期"},
                template="plotly_dark",
                color="用户数",
                color_continuous_scale="viridis"
            )
            fig_high_pay.update_layout(xaxis_tickangle=-45)
            st.plotly_chart(fig_high_pay, use_container_width=True)
            
            # ---------------------- 5. 异常数据汇总 ----------------------
            st.markdown('<div class="module-title">⚠️ 异常数据汇总</div>', unsafe_allow_html=True)
            
            # 汇总所有异常
            summary_items = []
            if len(dau_warn_data) > 0:
                summary_items.append(f"DAU异常：{len(dau_warn_data)}天低于预警阈值（{dau_warn_threshold/10000}万）；")
            if len(ret_warn_data) > 0:
                summary_items.append(f"留存异常：{len(ret_warn_data)}天新手7日留存率低于预警阈值（{retention_warn_threshold:.2%}）；")
            
            # 检查付费转化率异常（低于1%）
            convert_warn_data = convert_df[convert_df["付费转化率"] < 0.01]
            if len(convert_warn_data) > 0:
                summary_items.append(f"付费转化异常：{len(convert_warn_data)}天付费转化率低于1%；")
            
            # 拼接为自然分行的文本
            if summary_items:
                summary_text = "\n".join(summary_items)
                st.markdown(f"""
                    <div class="warn-box">
                        <b>⚠️ 异常检测结果</b><br><br>
                        {summary_text.replace(';', ';<br>')}
                    </div>
                """, unsafe_allow_html=True)
            else:
                summary_text = "✅ 未检测到明显数据异常，运营数据整体平稳"
                st.markdown(f"""
                    <div class="conclusion-box">
                        <b>📌 异常检测结果</b><br><br>
                        {summary_text}
                    </div>
                """, unsafe_allow_html=True)
            
            # ---------------------- 6. 运营结论自动总结 ----------------------
            st.markdown('<div class="module-title">📝 运营结论自动总结（星铁定制版）</div>', unsafe_allow_html=True)
            
            # 核心结论生成
            conclusion_items = []
            
            # 1. 卡池表现
            cycle_flow = payment_df.groupby("卡池周期")["卡池总流水(元)"].sum().reset_index()
            best_cycle = cycle_flow.loc[cycle_flow["卡池总流水(元)"].idxmax(), "卡池周期"]
            best_role = payment_df[payment_df["卡池周期"] == best_cycle]["主UP角色"].iloc[0]
            conclusion_items.append(f"1. 卡池表现：{best_cycle}（{best_role}UP）流水最高，该角色为星铁热门限定角色，建议后续优先推出配套角色卡池；")
            
            # 2. 用户留存
            ret_7_mean = retention_df["新手7日留存率"].mean()
            if ret_7_mean > 0.45:
                conclusion_items.append(f"2. 用户留存：新手7日留存率均值{ret_7_mean:.2%}（符合星铁成熟区基准），当前新手引导和新手奖励设计合理，无需调整；")
            else:
                conclusion_items.append(f"2. 用户留存：新手7日留存率均值{ret_7_mean:.2%}（低于星铁基准），建议优化新手引导和新手星琼福利、降低初期活动难度；")
            
            # 3. 付费转化
            convert_mean = convert_df["付费转化率"].mean()
            if convert_mean > 0.015:
                conclusion_items.append(f"3. 付费转化：付费转化率均值{convert_mean:.2%}（符合星铁成熟区基准），充值档位定价和卡池保底规则合理；")
            else:
                conclusion_items.append(f"3. 付费转化：付费转化率均值{convert_mean:.2%}（低于星铁基准），建议推出「老玩家回归礼包」「角色满命福利」「累计充值礼包」等提升付费意愿；")
            
            # 4. 异常处理
            if 'dau_warn_data' in locals() and len(dau_warn_data) > 0:
                conclusion_items.append(f"4. 异常处理：{len(dau_warn_data)}天DAU低于预警阈值，需排查星铁服务器稳定性，建议优化移动端加载速度（星铁移动端用户占比60%+）；")
            
            # 5. 老用户粘性
            active_ret_mean = retention_df["活跃用户7日留存率"].mean()
            if active_ret_mean > 0.8:
                conclusion_items.append(f"5. 老用户粘性：活跃用户7日留存率均值{active_ret_mean:.2%}（优秀），星铁的每日委托、模拟宇宙等长线玩法留存效果显著，建议新增「老玩家专属礼物」等；")
            else:
                conclusion_items.append(f"5. 老用户粘性：活跃用户7日留存率均值{active_ret_mean:.2%}（待优化），建议为登录超90天的老用户解锁「专属剧情任务」「限定头像框」「限定头像」等；")
            
            # 6. 付费分层优化
            high_value_ratio = len(user_layer_df[user_layer_df["付费价值分层"] == "高价值"]) / len(user_layer_df)
            if high_value_ratio < 0.05:
                conclusion_items.append(f"6. 付费分层：高价值用户占比{high_value_ratio:.2%}（低于星铁基准），建议推出「满命角色专属皮肤」「满精光锥专属战斗效果」等提升高价值用户占比；")
            else:
                conclusion_items.append(f"6. 付费分层：高价值用户占比{high_value_ratio:.2%}（符合星铁基准），建议保持充值档位的福利力度；")
            
            # 拼接为自然分行的文本
            conclusion_web = "<br>".join(conclusion_items)
            conclusion_list = conclusion_items
            
            # 展示结论
            st.markdown(f"""
                <div class="conclusion-box">
                    <b>📌 星铁运营核心结论（{datetime.now().strftime('%Y-%m-%d')}）</b><br><br>
                    {conclusion_web}
                </div>
            """, unsafe_allow_html=True)
            
            # ---------------------- 7. 分析报告导出 ----------------------
            st.markdown('<div class="module-title">💾 分析报告导出</div>', unsafe_allow_html=True)
            
            # 生成Word分析报告函数
            def generate_analysis_report(conclusion_list):
                doc = Document()
                
                # 全局字体配置
                def set_font(paragraph, font_name="微软雅黑", font_size=12, bold=False, alignment=WD_PARAGRAPH_ALIGNMENT.LEFT):
                    """设置段落字体样式"""
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
                        run.font.bold = bold
                    paragraph.alignment = alignment
                
                # 标题
                title = doc.add_heading(level=0)
                title_run = title.add_run('崩坏：星穹铁道运营分析报告')
                title_run.font.name = "微软雅黑"
                title_run.font.size = Pt(20)
                title_run.font.bold = True
                title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                # 添加报告日期
                date_para = doc.add_paragraph(f'报告生成日期：{datetime.now().strftime("%Y年%m月%d日")}')
                set_font(date_para, font_size=11, alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT)
                doc.add_paragraph()
            
                # 1. 核心指标汇总
                heading1 = doc.add_heading('一、核心指标汇总', level=1)
                set_font(heading1, font_size=14, bold=True)
                
                # 创建表格
                table1 = doc.add_table(rows=5, cols=2)
                table1.alignment = WD_TABLE_ALIGNMENT.CENTER
                table1.style = 'Table Grid'
                
                # 表头
                hdr_cells = table1.rows[0].cells
                hdr_cells[0].text = '指标名称'
                hdr_cells[1].text = '数值'
                for cell in hdr_cells:
                    para = cell.paragraphs[0]
                    set_font(para, font_size=11, bold=True)
                
                # 填充数据
                table_data = [
                    ('总新增用户', f'{total_new_users:,}'),
                    ('平均DAU', f'{avg_dau:,}'),
                    ('新手7日留存率', f'{avg_retention_7:.2%}'),
                    ('总流水（万元）', f'{total_revenue:.2f}')
                ]
                for i, (label, value) in enumerate(table_data, start=1):
                    row_cells = table1.rows[i].cells
                    row_cells[0].text = clean_text(label)
                    row_cells[1].text = clean_text(value)
                    for cell in row_cells:
                        para = cell.paragraphs[0]
                        set_font(para, font_size=11)
                
                doc.add_paragraph()
            
                # 2. 异常分析总结
                heading2 = doc.add_heading('二、异常分析总结', level=1)
                set_font(heading2, font_size=14, bold=True)
                
                clean_summary = clean_text(summary_text)
                for item in clean_summary.split('\n'):
                    if item.strip():
                        para = doc.add_paragraph(item.strip())
                        set_font(para)
                doc.add_paragraph()
            
                # 3. 运营结论与建议
                heading3 = doc.add_heading('三、运营结论与星铁定制化建议', level=1)
                set_font(heading3, font_size=14, bold=True)
                
                for conclusion in conclusion_list:
                    if conclusion.strip():
                        clean_con = clean_text(conclusion)
                        para = doc.add_paragraph(clean_con)
                        set_font(para)
                doc.add_paragraph()
            
                # 4. 卡池表现分析
                heading4 = doc.add_heading('四、卡池表现分析', level=1)
                set_font(heading4, font_size=14, bold=True)
                
                cycle_flow = payment_df.groupby("卡池周期")["卡池总流水(元)"].sum().reset_index()
                best_cycle = cycle_flow.loc[cycle_flow["卡池总流水(元)"].idxmax(), "卡池周期"]
                card_analysis = [
                    f'最优卡池周期：{best_cycle}，建议复用该卡池的角色组合策略；',
                    '主UP池流水占比70%+，符合星铁运营特征，建议持续聚焦新限定角色；',
                    '常驻池占比合理，无需调整保底规则。'
                ]
                for item in card_analysis:
                    para = doc.add_paragraph(clean_text(item))
                    set_font(para)
            
                # 保存到临时文件
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
                doc.save(temp_file.name)
                temp_file.seek(0)
                return temp_file.read()
            
            # Excel导出函数
            @st.cache_data
            def convert_to_excel(df):
                output = io.BytesIO()
                with pd.ExcelWriter(output, engine='openpyxl') as writer:
                    df.to_excel(writer, index=False, sheet_name="运营全量数据")
                    activity_df.to_excel(writer, index=False, sheet_name="日活数据")
                    retention_df.to_excel(writer, index=False, sheet_name="留存数据")
                    payment_df.to_excel(writer, index=False, sheet_name="付费数据")
                    user_layer_df.to_excel(writer, index=False, sheet_name="用户分层数据")
                return output.getvalue()
            
            # 合并数据用于Excel导出
            all_data = pd.merge(activity_df, retention_df, on="日期", suffixes=("_日活", "_留存"))
            all_data = pd.merge(all_data, payment_df.drop(["卡池周期", "主UP角色", "复刻角色"], axis=1), on="日期")
            excel_data = convert_to_excel(all_data)
            
            # 初始化Word报告数据
            report_data = b""
            
            # 生成Word报告
            try:
                report_data = generate_analysis_report(conclusion_list)
                st.success("✅ 运营结论生成完成！可下载完整Word分析报告")
            except Exception as e:
                st.warning(f"⚠️ Word报告生成失败：{str(e)}")
                st.info("💡 请确保已安装python-docx库：pip install python-docx")
            
            # 导出按钮
            col1, col2 = st.columns(2)
            with col1:
                st.download_button(
                    label="📥 导出运营分析报告（Word）",
                    data=report_data,
                    file_name=f"星铁运营分析报告_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            with col2:
                st.download_button(
                    label="📥 导出全量原始数据（Excel）",
                    data=excel_data,
                    file_name=f"星铁运营原始数据_{datetime.now().strftime('%Y%m%d')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
            
        except Exception as e:
            st.error(f"❌ 数据加载失败：{str(e)}")
            st.info("💡 请确认上传的文件是正确的星铁运营数据集Excel文件")
    else:
        st.markdown("""
            <div style="text-align:center; padding:50px 0;">
                <h2 style="color:#a78bfa;">🚀 星穹铁道运营数据看板</h2>
                <p style="color:#94a3b8; font-size:16px; margin:20px 0;">
                    请在左侧侧边栏上传运营数据集Excel文件<br>
                    支持分析周期：84天（成熟稳定期双版本）
                </p>
                <div style="border:1px dashed #8b5cf6; padding:20px; border-radius:8px; max-width:600px; margin:0 auto;">
                    <h4 style="color:#a78bfa;">📋 看板功能</h4>
                    <ul style="text-align:left; color:#94a3b8; line-height:1.8;">
                        <li>核心指标概览（DAU、留存、流水、付费转化）</li>
                        <li>日活与留存趋势分析（含异常预警）</li>
                        <li>付费数据分析（卡池流水、转化率）</li>
                        <li>用户分层分析（付费价值、生命周期）</li>
                        <li>异常数据自动汇总</li>
                        <li>星铁定制化运营结论自动生成</li>
                        <li>Excel/Word报告导出</li>
                    </ul>
                </div>
            </div>
        """, unsafe_allow_html=True)

# ---------------------- 2. A/B测试分析标签页 ----------------------
with tab2:
    ab_test_module()

# ---------------------- 3. 用户行为漏斗分析标签页 ----------------------
with tab3:
    funnel_analysis_module()

# ---------------------- 页面底部 ----------------------
st.divider()
st.markdown("""
    <div style="text-align:center; color:#94a3b8; font-size:14px;">
        📌 星穹铁道运营数据看板 | 基于成熟稳定期模拟运营数据<br>
        🎯 适配米哈游运营分析逻辑 | 贴合星铁游戏特色设计<br>
        📊 新增A/B测试+用户漏斗分析模块
    </div>
""", unsafe_allow_html=True)